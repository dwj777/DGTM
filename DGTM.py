import time
import re
import itertools
from pathlib import Path
import numpy as np
from scipy.optimize import fsolve
from scipy.linalg import eigvals
from scipy.integrate import solve_ivp
from docx import Document

_eig_cache = {}

DUP_TOL = 1e-6
FINAL_DUP_TOL = 1e-3
CAT_HIGH_EPS = 1e-3
CAT_M_LO = 0.2
def make_incremented_outpath(base_out_path, run_offset=0):
    p = Path(base_out_path)
    parent = p.parent if str(p.parent) != '.' else Path.cwd()
    stem = p.stem
    suffix = p.suffix or ".docx"
    m = re.search(r'(.*?)([0-9]+)$', stem)
    if m:
        prefix = m.group(1)
        base_num = int(m.group(2))
        target_num = base_num + run_offset
        new_stem = f"{prefix}{target_num}"
    else:
        if run_offset == 0:
            new_stem = stem
        else:
            new_stem = f"{stem}_{run_offset}"
    candidate = parent / (new_stem + suffix)
    if candidate.exists():
        if m:
            start = base_num + max(0, run_offset)
            i = start
            while True:
                cand = parent / (f"{prefix}{i}" + suffix)
                if not cand.exists():
                    return str(cand)
                i += 1
        else:
            i = run_offset if run_offset > 0 else 1
            while True:
                cand = parent / (f"{stem}_{i}" + suffix)
                if not cand.exists():
                    return str(cand)
                i += 1
    else:
        return str(candidate)

#  (1)Parameter sampling
def lhs(n_samples, n_dims, random_state=None, rng=None):
    if rng is None:
        rng = np.random.default_rng(random_state)
    u = rng.uniform(size=(n_samples, n_dims))
    points = np.zeros((n_samples, n_dims))
    for j in range(n_dims):
        perm = rng.permutation(n_samples)
        points[:, j] = (perm + u[:, j]) / n_samples
    return points

def model(y, t, params):
    if isinstance(params, dict):
        merged = {}
        merged.update(params)
        if 'weights' in params and isinstance(params['weights'], dict):
            merged.update(params['weights'])
        for sub in ('P', 'G', 'T', 'M'):
            if sub in params and isinstance(params[sub], dict):
                merged.update(params[sub])
        params = merged

    P, G, T, M = y
    alpha = params["alpha"]
    gamma = params["gamma"]
    K_MtoP = params["K_MtoP"]; hm2p = params.get("hm2p", 3.0)
    K_GtoP = params["K_GtoP"]; hg2p = params.get("hg2p", 3.0)
    beta_P = params["beta_P"]; K_Pauto = params["K_Pauto"]; hp2p = params.get("hp2p", 3.0)
    K_PtoG = params["K_PtoG"]; hp2g = params.get("hp2g", 3.0)
    K_TtoG = params["K_TtoG"]; ht2g = params.get("ht2g", 3.0)
    K_MtoG = params["K_MtoG"]; hm2g = params.get("hm2g", 3.0)
    K_GtoT = params["K_GtoT"]; hg2t = params.get("hg2t", 3.0)
    K_Tauto = params["K_Tauto"]; ht2t = params.get("ht2t", 3.0)
    K_GtoM = params["K_GtoM"]; hg2m = params.get("hg2m", 3.0)
    K_TtoM = params["K_TtoM"]; ht2m = params.get("ht2m", 3.0)
    K_PtoM = params["K_PtoM"]; hp2m = params.get("hp2m", 3.0)
    w_P = params["w_P"]; w_G = params["w_G"]; w_T = params["w_T"]

    basal_P = 0.01; basal_G = 0.01; basal_T = 0.03; basal_M = 0.05

    fG2P = (1 + 1 / (1 + (G / K_GtoP) ** hg2p))
    dPdt = (basal_P
            + alpha * (1 / (1 + (M / K_MtoP) ** hm2p)) * fG2P
            + beta_P * (P ** hp2p / (K_Pauto ** hp2p + P ** hp2p))
           ) - gamma * P

    dGdt = (basal_G
            + alpha * (P ** hp2g / (K_PtoG ** hp2g + P ** hp2g))
                     * (T ** ht2g / (K_TtoG ** ht2g + T ** ht2g))
                     * (1 / (1 + (M / K_MtoG) ** hm2g))
            - gamma * G)

    dTdt = (basal_T
            + alpha * (G ** hg2t / (K_GtoT ** hg2t + G ** hg2t))
                      * (1 / (1 + (T / K_Tauto) ** ht2t))
            - gamma * T)

    dMdt = (basal_M
            + alpha * (1 / (1 + w_P * (P / K_PtoM) ** hp2m + w_G * (G / K_GtoM) ** hg2m + w_T * (T / K_TtoM) ** ht2m))
            - gamma * M)

    return np.array([dPdt, dGdt, dTdt, dMdt], dtype=float)

def sample_parameters_from_u(u):
    def rnd(x): return round(float(x), 6)
    def map_lin(uval, a, b): return a + uval * (b - a)
    def map_log(uval, a, b): return 10 ** (np.log10(a) + uval * (np.log10(b) - np.log10(a)))

    n_dims = 17
    if len(u) != n_dims:
        raise ValueError(f"u must have length {n_dims}, got {len(u)}")

    alpha = rnd(map_lin(u[0], 0.1, 3.0))
    beta_P = rnd(map_lin(u[1], 0.1, 3.0))
    gamma = rnd(map_lin(u[2], 0.01, 0.5))

    K_MtoP = rnd(map_log(u[3], 0.1, 5.0))
    K_GtoP = rnd(map_log(u[4], 0.1, 5.0))
    K_Pauto = rnd(map_log(u[5], 0.1, 5.0))

    K_PtoG = rnd(map_log(u[6], 0.1, 5.0))
    K_TtoG = rnd(map_log(u[7], 0.1,5.0))
    K_MtoG = rnd(map_log(u[8], 0.1, 5.0))

    K_GtoT = rnd(map_log(u[9], 0.1,5.0))
    K_Tauto = rnd(map_log(u[10], 0.1, 5.0))

    K_GtoM = rnd(map_log(u[11], 0.1, 5.0))
    K_TtoM = rnd(map_log(u[12], 0.1, 5.0))

    K_PtoM = rnd(map_log(u[13], 0.1, 5.0))

    w_P = rnd(map_lin(u[14], 0.1, 1.0))
    w_G = rnd(map_lin(u[15], 0.1, 1.0))
    w_T = rnd(map_lin(u[16], 0.1, 1.0))

    h_fixed = 3.0

    pP = {'K_MtoP': K_MtoP, 'hm2p': h_fixed, 'K_GtoP': K_GtoP, 'hg2p': h_fixed,
          'K_Pauto': K_Pauto, 'hp2p': h_fixed, 'beta_P': beta_P}

    pG = {'K_PtoG': K_PtoG, 'hp2g': h_fixed, 'K_TtoG': K_TtoG, 'ht2g': h_fixed,
          'K_MtoG': K_MtoG, 'hm2g': h_fixed}

    pT = {'K_GtoT': K_GtoT, 'hg2t': h_fixed,
          'K_Tauto': K_Tauto, 'ht2t': h_fixed}

    pM = {'K_GtoM': K_GtoM, 'hg2m': h_fixed, 'K_TtoM': K_TtoM, 'ht2m': h_fixed,
          'K_PtoM': K_PtoM, 'hp2m': h_fixed}

    p_weights = {'w_P': w_P, 'w_G': w_G, 'w_T': w_T}
    p_common = {'alpha': alpha, 'gamma': gamma}
    nested = {'P': pP, 'G': pG, 'T': pT, 'M': pM, 'weights': p_weights}
    flat = {}
    flat.update(p_common); flat.update(nested)
    for d in (pP, pG, pT, pM, p_weights): flat.update(d)
    return flat

def dedupe_points(pts, tol=DUP_TOL):
    pts_arr = np.asarray(pts, dtype=float)
    unique = []
    for p in pts_arr:
        is_dup = False
        for u in unique:
            if np.linalg.norm(p - u) <= tol:
                is_dup = True; break
        if not is_dup: unique.append(p.copy())
    return unique

# (2)Equilibrium searching
def solve_steady_states_lhs(params, n_inits=1000, bounds=((0.0,10.0),)*4, tol_res=1e-6, rng=None, fsolve_maxfev=2000, xtol=1e-8):
    if rng is None: rng = np.random.default_rng()
    n_dims = 4
    lhs_pts = lhs(n_inits, n_dims, rng=rng)
    D_bounds, G_bounds, T_bounds, M_bounds = bounds
    initials = np.column_stack([
        D_bounds[0] + lhs_pts[:,0] * (D_bounds[1] - D_bounds[0]),
        G_bounds[0] + lhs_pts[:,1] * (G_bounds[1] - G_bounds[0]),
        T_bounds[0] + lhs_pts[:,2] * (T_bounds[1] - T_bounds[0]),
        M_bounds[0] + lhs_pts[:,3] * (M_bounds[1] - M_bounds[0])
    ])

    sols = []
    for x0 in initials:
        try:
            sol, infodict, ier, mesg = fsolve(lambda x: model(x, 0, params), x0, full_output=True, xtol=xtol, maxfev=fsolve_maxfev)
            if ier != 1:
                continue
            sol = np.real(sol)
            fval = model(sol, 0, params)
            if not np.all(np.isfinite(fval)): continue
            res_norm = np.linalg.norm(fval)
            if (not np.isfinite(res_norm)) or (res_norm > tol_res):
                continue
            # bounds check
            if not (D_bounds[0] <= sol[0] <= D_bounds[1] and G_bounds[0] <= sol[1] <= G_bounds[1] and T_bounds[0] <= sol[2] <= T_bounds[1] and M_bounds[0] <= sol[3] <= M_bounds[1]):
                continue
            sols.append(sol)
        except Exception:
            continue

    uniq = []
    tol = 1e-6
    for s in sols:
        if not any(np.linalg.norm(s - u) <= tol for u in uniq):
            uniq.append(s)
    return uniq

# (3) Stability assessment
def jacobian_and_score(pt, params, base_eps=1e-6):
    pt_arr = np.asarray(pt, dtype=float)
    key = (id(params), tuple(np.round(pt_arr, 6)))
    if key in _eig_cache:
        return _eig_cache[key]
    n = len(pt_arr)
    J = np.zeros((n, n), dtype=float)
    for i in range(n):
        eps_i = base_eps * max(1.0, abs(pt_arr[i]))
        dp = np.zeros(n, dtype=float); dp[i] = eps_i
        f_plus = model(pt_arr + dp, 0, params)
        f_minus = model(pt_arr - dp, 0, params)
        try:
            f_plus = np.asarray(f_plus, dtype=float); f_minus = np.asarray(f_minus, dtype=float)
        except Exception:
            return None, None, -np.inf
        if f_plus.shape != (n,) or f_minus.shape != (n,): return None, None, -np.inf
        if not (np.all(np.isfinite(f_plus)) and np.all(np.isfinite(f_minus))): return None, None, -np.inf
        J[:, i] = (f_plus - f_minus) / (2.0 * eps_i)
    if not np.all(np.isfinite(J)): return None, None, -np.inf
    try:
        vals = eigvals(J)
    except Exception:
        return None, None, -np.inf
    if not np.all(np.isfinite(vals)): return None, None, -np.inf
    score = -np.max(np.real(vals))
    _eig_cache[key] = (J, vals, score)
    return J, vals, score

def is_stable_and_score(pt, params):
    J, vals, score = jacobian_and_score(pt, params)
    if J is None or vals is None or score == -np.inf:
        return False, score
    return np.all(np.real(vals) < 0), score

# (4) Trajectory convergence verification
def integration_stability_check(pt, params, rng=None, perturb_rel=1e-3, int_t_max=50, int_atol=1e-6, int_rtol=1e-4, final_tol=1e-2):
    if rng is None: rng = np.random.default_rng()
    delta = np.maximum(np.abs(pt) * perturb_rel, perturb_rel)
    y0 = pt + rng.uniform(-1, 1, size=pt.shape) * delta
    try:
        sol = solve_ivp(lambda t, y: model(y, t, params), [0, int_t_max], y0, method='BDF', atol=int_atol, rtol=int_rtol)
        if not getattr(sol, 'success', False): return False
        if not np.all(np.isfinite(sol.y)): return False
        y_end = sol.y[:, -1]
        err = np.linalg.norm(y_end - pt)
        if not np.isfinite(err): return False
        tol = final_tol * (1.0 + np.linalg.norm(pt))
        return err <= tol
    except Exception:
        return False

# (5) Pattern matching
def categorize_value(v, eps_high=CAT_HIGH_EPS, m_lo=CAT_M_LO):

    try:
        if not np.isfinite(v):
            return None
    except Exception:
        return None

    if v >= 1.0 - eps_high:
        return 'H'
    if v > m_lo and v < 1.0 - eps_high:
        return 'M'
    return 'L'

def passes_norm_pattern_row(normalized_matrix, debug=False):

    arr = np.asarray(normalized_matrix, dtype=float)
    if arr.shape != (3,4):
        if debug:
            print("[passes_norm_pattern_row] invalid shape:", arr.shape)
        return False
    if not np.all(np.isfinite(arr)):
        if debug:
            print("[passes_norm_pattern_row] non-finite in arr")
        return False

    cats = np.empty(arr.shape, dtype=object)
    for i in range(3):
        for j in range(4):
            cats[i, j] = categorize_value(arr[i, j])
            if cats[i, j] is None:
                if debug:
                    print(f"[passes_norm_pattern_row] None at {i},{j}")
                return False

    def is_H(x): return x == 'H'
    def is_M(x): return x == 'M'
    def is_L(x): return x == 'L'
    def is_M_or_L(x): return (x == 'M' or x == 'L')

    for perm in itertools.permutations(range(3)):
        rA = cats[perm[0]]
        rB = cats[perm[1]]
        rC = cats[perm[2]]


        condA = (is_H(rA[0]) and is_M_or_L(rA[1]) and is_M_or_L(rA[2]) and is_M_or_L(rA[3]))

        condC = (is_L(rC[0]) and is_L(rC[1]) and is_L(rC[2]) and is_H(rC[3]))

        if not (condA and condC):
            continue

        condB = (is_M(rB[0]) and is_M(rB[1]) and is_M(rB[2]) and is_H(rB[3]))

        if condB:
            if debug:
                print("[passes_norm_pattern_row] matched strict A/B/C with perm", perm, "cats:", cats)
            return True

    if debug:
        print("[passes_norm_pattern_row] no match. cats:", cats)
    return False

def save_results_to_word(accepted, out_path):
    try:
        doc = Document()
        doc.add_heading('DGTM Model (LHS + fsolve) Results', level=1)
        for i, item in enumerate(accepted, 1):
            if len(item) == 3:
                p, reps_with_scores, normalized = item
            else:
                p, reps_with_scores = item
                normalized = None
            doc.add_heading(f'Parameter Set #{i}', level=2)
            for f in ['P','G','T','M']:
                if f in p and isinstance(p[f], dict):
                    items = sorted(p[f].items(), key=lambda x: x[0])
                    doc.add_paragraph(f"{f} parameters: " + ", ".join(f"{k}={v:.6g}" for k,v in items))
            if 'weights' in p and isinstance(p['weights'], dict):
                items = sorted(p['weights'].items(), key=lambda x: x[0])
                doc.add_paragraph("weights parameters: " + ", ".join(f"{k}={v:.6g}" for k,v in items))
            excluded = set(['P','G','T','M','weights'])
            top_level_keys = sorted([k for k,v in p.items() if not isinstance(v, dict) and k not in excluded])
            if top_level_keys:
                doc.add_paragraph("Other top-level parameters: " + ", ".join(f"{k}={p[k]:.6g}" for k in top_level_keys))
            doc.add_paragraph(f"Number of steady states: {len(reps_with_scores)}")
            if normalized is not None and len(normalized) == len(reps_with_scores):
                doc.add_paragraph("(Each row: original steady state P,G,T,M | corresponding row-normalized P,G,T,M)")
                for (pt, score), norm_row in zip(reps_with_scores, normalized):
                    orig_str = ", ".join(f"{name}={val:.6f}" for name,val in zip(['P','G','T','M'], pt))
                    norm_str = ", ".join(f"{v:.6f}" for v in norm_row)
                    doc.add_paragraph(f"  {orig_str}  |  normalized_row: {norm_str}  |  jacobian_score={score:.6f}")
            else:
                for pt, score in reps_with_scores:
                    doc.add_paragraph("  " + ", ".join(f"{name}={val:.6f}" for name,val in zip(['P','G','T','M'], pt)) + f"  | jacobian_score={score:.6f}")
            doc.add_paragraph()
        doc.save(out_path)
    except Exception as e:
        print(f"[save] Error saving Word: {e}")

def main(random_state=None,
         n_param_samples=10000,
         n_inits_per_param=1000,
         fsolve_maxfev=2000,
         tol_root_res=1e-6,
         bounds=((0.0,10.0),(0.0,10.0),(0.0,10.0),(0.0,10.0)),
         out_path=r"DGTM_selected_params_lhs_fsolve.docx",
         max_iterations=None):
    start_time = time.time()
    rng = np.random.default_rng(random_state)

    n_dims = 17
    U_pool = lhs(n_param_samples, n_dims, rng=rng)
    U_list = [tuple(u) for u in U_pool]
    if max_iterations is not None:
        U_list = U_list[:max_iterations]
    print(f"[main] total parameter candidates to evaluate: {len(U_list)}")
    counters = {'iterations':0, 'accepted':0}
    accepted = []
    for idx, u_tup in enumerate(U_list, start=1):
        counters['iterations'] = idx
        u = np.array(u_tup)
        params = sample_parameters_from_u(u)
        _eig_cache.clear()
        sols = solve_steady_states_lhs(params, n_inits=n_inits_per_param, bounds=bounds, tol_res=tol_root_res, rng=rng, fsolve_maxfev=fsolve_maxfev)
        if not sols:
            continue
        uniq_roots = dedupe_points(sols, tol=1e-6)
        if len(uniq_roots) < 3:
            continue
        jacobian_passed_pts = []
        jacobian_scores = []
        for pt in uniq_roots:
            stable_jac, score = is_stable_and_score(pt, params)
            if not stable_jac: continue
            jacobian_passed_pts.append(pt); jacobian_scores.append(score)
        if len(jacobian_passed_pts) < 3:
            continue
        stable_after_both = []
        stable_scores = []
        for pt, sc in zip(jacobian_passed_pts, jacobian_scores):
            ok = integration_stability_check(pt, params, rng=rng, int_t_max=50, final_tol=1e-2)
            if not ok: continue
            stable_after_both.append(pt); stable_scores.append(sc)
        atol_dup = FINAL_DUP_TOL
        unique_pts = []
        unique_scores = []
        for pt, sc in zip(stable_after_both, stable_scores):
            is_dup = False
            for up in unique_pts:
                if np.all(np.abs(pt - up) <= atol_dup): is_dup = True; break
            if not is_dup:
                unique_pts.append(np.asarray(pt, dtype=float)); unique_scores.append(float(sc))
        n_unique = len(unique_pts)
        if n_unique != 3:
            continue
        reps_final = [(unique_pts[i], unique_scores[i]) for i in range(n_unique)]
        arr = np.vstack(unique_pts)
        row_max = np.max(np.abs(arr), axis=1)
        row_max[row_max == 0.0] = 1.0
        normalized_matrix = (arr / row_max[:, None])
        if not passes_norm_pattern_row(normalized_matrix):
            continue
        normalized_list = normalized_matrix.tolist()
        accepted.append((params, reps_final, normalized_list)); counters['accepted'] += 1
        print(f"=== Found accepted #{counters['accepted']} at candidate #{idx}/{len(U_list)} ===")
        print("Format: original P,G,T,M | row-normalized P,G,T,M | jacobian_score")
        for (pt, sc), norm_row in zip(reps_final, normalized_list):
            orig_str = ", ".join(f"{v:.6f}" for v in pt)
            norm_str = ", ".join(f"{v:.6f}" for v in norm_row)
            print(f"  {orig_str}  |  {norm_str}  |  score={sc:.6f}")
    total_time = time.time() - start_time
    print("=== SEARCH FINISHED ===")
    print(f"Total candidates processed: {counters['iterations']}")
    print(f"accepted: {counters['accepted']}")
    print(f"Total elapsed time: {total_time:.1f}s")
    try:
        save_results_to_word(accepted, out_path)
        print(f"[final save] saved {len(accepted)} accepted sets to {out_path}")
    except Exception as e:
        print(f"[final save] Error saving Word: {e}")
    return accepted

if __name__ == "__main__":
    BASE_OUT_PATH = r"C:\Users\asus\Desktop\DGTM1.docx"
    NUM_RUNS = 5
    for run_idx in range(NUM_RUNS):
        cur_out = make_incremented_outpath(BASE_OUT_PATH, run_offset=run_idx)
        print(f"==== START RUN {run_idx+1}/{NUM_RUNS} -> saving to: {cur_out} ====")
        out = main(
            random_state=None,
            n_param_samples=20000,
            n_inits_per_param=800,
            fsolve_maxfev=1000,
            tol_root_res=1e-6,
            bounds=((0.0,10.0),(0.0,10.0),(0.0,10.0),(0.0,10.0)),
            out_path=cur_out,
            max_iterations=None
        )
        print(f"Run {run_idx+1} finished. Accepted count: {0 if out is None else len(out)}")
    print("All runs completed.")
